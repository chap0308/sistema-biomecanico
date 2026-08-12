from __future__ import annotations

import argparse
import subprocess
import tempfile
from copy import deepcopy
from pathlib import Path

import markdown
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")


def build_html(markdown_path: Path) -> str:
    source = markdown_path.read_text(encoding="utf-8")
    body = markdown.markdown(
        source,
        extensions=["tables", "fenced_code", "sane_lists"],
        output_format="html5",
    )
    return f"""<!doctype html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>{markdown_path.stem}</title>
<style>
  @page {{ size: A4; margin: 2.5cm; }}
  body {{ font-family: "Times New Roman", serif; font-size: 11pt; line-height: 2; color: #111; }}
  h1 {{ font-size: 14pt; margin: 18pt 0 10pt; page-break-after: avoid; }}
  h2 {{ font-size: 12pt; margin: 14pt 0 8pt; page-break-after: avoid; }}
  h3 {{ font-size: 11pt; margin: 12pt 0 6pt; page-break-after: avoid; }}
  p {{ margin: 0 0 8pt; text-align: justify; }}
  li {{ margin-bottom: 4pt; }}
  table {{ border-collapse: collapse; width: 100%; margin: 10pt 0 16pt; font-size: 8pt; line-height: 1.15; page-break-inside: auto; }}
  th, td {{ border: 0.6pt solid #666; padding: 4pt; vertical-align: middle; text-align: left; overflow-wrap: anywhere; }}
  th {{ background: #d9e8f5; font-weight: bold; }}
  code {{ font-family: Consolas, monospace; font-size: 9pt; }}
  a {{ color: #111; text-decoration: none; }}
</style>
</head>
<body>{body}</body>
</html>"""


def convert(markdown_path: Path, output_path: Path) -> None:
    if not SOFFICE.exists():
        raise FileNotFoundError(f"No se encontró LibreOffice en {SOFFICE}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="thesis-md-docx-") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        html_path = temp_dir / f"{output_path.stem}.html"
        html_path.write_text(build_html(markdown_path), encoding="utf-8")
        subprocess.run(
            [
                str(SOFFICE),
                "--headless",
                "--convert-to",
                "docx:Office Open XML Text",
                "--outdir",
                str(temp_dir),
                str(html_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        generated = temp_dir / f"{output_path.stem}.docx"
        if not generated.exists():
            raise RuntimeError("LibreOffice no generó el archivo DOCX esperado")
        output_path.write_bytes(generated.read_bytes())
    apply_annex_landscape_section(output_path)


def apply_annex_landscape_section(output_path: Path) -> None:
    document = Document(output_path)
    for table in document.tables:
        table.autofit = True
        table_properties = table._tbl.tblPr
        borders = table_properties.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            table_properties.append(borders)
        for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge = borders.find(qn(f"w:{edge_name}"))
            if edge is None:
                edge = OxmlElement(f"w:{edge_name}")
                borders.append(edge)
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "4")
            edge.set(qn("w:color"), "808080")
        for row in table.rows:
            row_properties = row._tr.get_or_add_trPr()
            for header_property in row_properties.findall(qn("w:tblHeader")):
                row_properties.remove(header_property)

    annex_heading = next(
        (paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "12. ANEXOS"),
        None,
    )
    if annex_heading is None:
        return

    portrait_properties = deepcopy(document.sections[-1]._sectPr)
    section_break = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    paragraph_properties.append(portrait_properties)
    section_break.append(paragraph_properties)
    annex_heading._p.addprevious(section_break)

    annex_section = document.sections[-1]
    annex_section.orientation = WD_ORIENT.LANDSCAPE
    annex_section.page_width = Cm(29.7)
    annex_section.page_height = Cm(21)
    annex_section.top_margin = Cm(2)
    annex_section.bottom_margin = Cm(2)
    annex_section.left_margin = Cm(1.5)
    annex_section.right_margin = Cm(1.5)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convierte el snapshot Markdown de tesis a DOCX.")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    convert(args.input.resolve(), args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
